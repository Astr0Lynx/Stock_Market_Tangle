#!/usr/bin/env python3
"""Convert REPORT.md to a formatted DOCX document."""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

import re

def add_table_from_markdown(doc, table_text):
    """Parse markdown table and add to document."""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return
    
    # Parse header
    header = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    num_cols = len(header)
    
    # Parse rows (skip separator line)
    rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if len(cells) == num_cols:
            rows.append(cells)
    
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Add header
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Add data rows
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    return table

def add_code_block(doc, code_text, language='python'):
    """Add formatted code block."""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    
    # Add light gray background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading_elm)
    
    return p

def process_markdown_formatting(paragraph, text):
    """Process bold, italic, and inline code in text."""
    # Handle **bold**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif '`' in part:
            # Handle inline code
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                elif cp:
                    paragraph.add_run(cp)
        elif part:
            paragraph.add_run(part)

def main():
    # Read markdown file
    with open('REPORT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block = []
            else:
                in_code_block = False
                if code_block:
                    add_code_block(doc, '\n'.join(code_block))
                    code_block = []
            i += 1
            continue
        
        if in_code_block:
            code_block.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            add_table_from_markdown(doc, '\n'.join(table_lines))
            table_lines = []
            doc.add_paragraph()
            continue
        
        # Title (first line)
        if i == 0 and not line.startswith('#'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True
            i += 1
            continue
        
        # Main title
        if line.strip() == 'Stock Market Tangle: A Graph-Theoretic Model of Financial Markets':
            doc.add_page_break()
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            i += 1
            continue
        
        # Section headings
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            i += 1
            continue
        elif line.startswith('# '):
            p = doc.add_paragraph(line[2:])
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            i += 1
            continue
        
        # Numbered sections (4.1, 5.2, etc.)
        if re.match(r'^\d+\.\d+', line.strip()):
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(13)
                run.font.bold = True
            i += 1
            continue
        
        # Major sections (Abstract, Introduction, etc.)
        if line.strip() in ['Abstract', 'Introduction', 'Algorithm Descriptions', 'Implementation Details',
                           'Experimental Setup', 'Results & Analysis', 'Conclusion', 'Bonus Disclosure', 'References']:
            doc.add_paragraph()
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 70, 127)
            i += 1
            continue
        
        # Bullet points
        if line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            process_markdown_formatting(p, line.strip()[2:])
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            process_markdown_formatting(p, text)
            i += 1
            continue
        
        # Horizontal rules
        if line.strip() in ['---', '___', '***']:
            doc.add_paragraph()
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraphs
        p = doc.add_paragraph()
        process_markdown_formatting(p, line.strip())
        i += 1
    
    # Save document
    output_file = 'REPORT.docx'
    doc.save(output_file)
    print(f'✓ Created {output_file}')

if __name__ == '__main__':
    main()
